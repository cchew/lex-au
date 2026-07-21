import pytest
from pathlib import Path
from docx import Document
from docx.oxml import parse_xml
from lexau.docx_reader import iter_paragraphs
from lexau.parser import ElementType

CORPUS_DOCX = Path(__file__).parent.parent / "corpus" / "docx"


def test_paragraph_passthrough():
    doc = Document()
    doc.add_paragraph("Hello world")
    results = list(iter_paragraphs(doc))
    texts = [p.text for p in results if p.text.strip()]
    assert "Hello world" in texts


def test_table_extracted_as_table_type():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "H1"
    table.cell(0, 1).text = "H2"
    table.cell(1, 0).text = "D1"
    table.cell(1, 1).text = "D2"
    results = list(iter_paragraphs(doc))
    table_blocks = [p for p in results if p.element_type == ElementType.TABLE]
    assert len(table_blocks) == 1
    assert table_blocks[0].table_rows == [["H1", "H2"], ["D1", "D2"]]


def test_document_order_preserved():
    doc = Document()
    doc.add_paragraph("Before")
    t = doc.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "Cell"
    doc.add_paragraph("After")
    results = list(iter_paragraphs(doc))
    meaningful = [p for p in results if p.text.strip() or p.element_type == ElementType.TABLE]
    table_idx = next(i for i, p in enumerate(meaningful) if p.element_type == ElementType.TABLE)
    assert meaningful[table_idx - 1].text.strip() == "Before"
    assert meaningful[table_idx + 1].text.strip() == "After"


def test_empty_paragraph_yields_skip():
    doc = Document()
    doc.add_paragraph("")
    results = list(iter_paragraphs(doc))
    skip_results = [p for p in results if p.element_type == ElementType.SKIP]
    assert len(skip_results) >= 1


def test_list_level_detection():
    """A DOCX paragraph with numPr at ilvl=0 yields LIST_ITEM with number='0'."""
    doc = Document()
    p = doc.add_paragraph("Plain list item text")
    # Inject numPr to mark paragraph as a list item at level 0
    pPr_xml = (
        '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:numPr>'
        '<w:ilvl w:val="0"/>'
        '<w:numId w:val="1"/>'
        '</w:numPr>'
        '</w:pPr>'
    )
    p._element.insert(0, parse_xml(pPr_xml))

    results = list(iter_paragraphs(doc))
    list_items = [r for r in results if r.element_type == ElementType.LIST_ITEM]
    assert len(list_items) == 1
    assert list_items[0].number == "0"
    assert list_items[0].text == "Plain list item text"


def test_list_level_detection_nested():
    """A list paragraph at ilvl=1 yields LIST_ITEM with number='1'."""
    doc = Document()
    p = doc.add_paragraph("Nested list item")
    pPr_xml = (
        '<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:numPr>'
        '<w:ilvl w:val="1"/>'
        '<w:numId w:val="1"/>'
        '</w:numPr>'
        '</w:pPr>'
    )
    p._element.insert(0, parse_xml(pPr_xml))

    results = list(iter_paragraphs(doc))
    list_items = [r for r in results if r.element_type == ElementType.LIST_ITEM]
    assert len(list_items) == 1
    assert list_items[0].number == "1"


def test_figure_element_type():
    """A paragraph containing an a:blip DrawingML element yields FIGURE element type."""
    doc = Document()
    p = doc.add_paragraph("Figure caption text")
    # Inject a DrawingML inline image blip into the paragraph element
    blip_xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:drawing>'
        '<wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">'
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<a:blip r:embed="rId1" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
        '</a:graphicData>'
        '</a:graphic>'
        '</wp:inline>'
        '</w:drawing>'
        '</w:r>'
    )
    p._element.append(parse_xml(blip_xml))

    results = list(iter_paragraphs(doc))
    figures = [r for r in results if r.element_type == ElementType.FIGURE]
    assert len(figures) == 1
    assert figures[0].text == "Figure caption text"


from lexau.parser import InlineSpan


def test_italic_run_populates_spans():
    """A paragraph with one italic run yields ParsedParagraph with spans."""
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("personal information")
    run.italic = True
    results = list(iter_paragraphs(doc))
    body = [r for r in results if r.text.strip() == "personal information"]
    assert len(body) == 1
    pp = body[0]
    assert len(pp.spans) == 1
    assert pp.spans[0].text == "personal information"
    assert pp.spans[0].italic is True
    assert pp.spans[0].bold is False


def test_bold_run_populates_spans():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("important")
    run.bold = True
    results = list(iter_paragraphs(doc))
    body = [r for r in results if r.text.strip() == "important"]
    assert len(body) == 1
    assert body[0].spans[0].bold is True


def test_mixed_runs_populate_spans():
    """Italic run followed by plain run yields two spans."""
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("term")
    r1.italic = True
    r2 = p.add_run(" means something")
    results = list(iter_paragraphs(doc))
    body = [r for r in results if "term" in r.text]
    assert len(body) == 1
    pp = body[0]
    assert len(pp.spans) == 2
    assert pp.spans[0].italic is True
    assert pp.spans[0].text == "term"
    assert pp.spans[1].italic is False
    assert pp.spans[1].text == " means something"


def test_plain_runs_span_not_formatted():
    doc = Document()
    p = doc.add_paragraph("hello world")
    results = list(iter_paragraphs(doc))
    body = [r for r in results if r.text.strip() == "hello world"]
    assert len(body) == 1
    pp = body[0]
    # spans may be populated but none should be formatted
    assert all(not (s.bold or s.italic or s.superscript or s.subscript) for s in pp.spans)


def test_superscript_run_populates_spans():
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("2")
    r.font.superscript = True
    results = list(iter_paragraphs(doc))
    body = [r for r in results if r.text.strip() == "2"]
    assert len(body) == 1
    assert body[0].spans[0].superscript is True


def test_image_paragraph_spans_empty():
    """FIGURE paragraphs (inline image) still have empty spans."""
    doc = Document()
    p = doc.add_paragraph("Figure caption")
    blip_xml = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:drawing>'
        '<wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">'
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<a:blip r:embed="rId1" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
        '</a:graphicData>'
        '</a:graphic>'
        '</wp:inline>'
        '</w:drawing>'
        '</w:r>'
    )
    p._element.append(parse_xml(blip_xml))
    results = list(iter_paragraphs(doc))
    figures = [r for r in results if r.element_type.value == "figure"]
    assert len(figures) == 1
    assert figures[0].spans == []


def test_loan_act_1976_shape1_fixture():
    # Shape 1: separate bold heading + single-tab numbered body.
    # Expect 5 sections (Short title, Commencement, Authority to borrow,
    # Application of moneys borrowed, Expenses of borrowing), 0 subsections.
    doc = Document(str(CORPUS_DOCX / "loan-act-(no.-2)-1976-vol0.docx"))
    paras = list(iter_paragraphs(doc))
    sections = [p for p in paras if p.element_type == ElementType.SECTION]
    subsections = [p for p in paras if p.element_type == ElementType.SUBSECTION]
    assert len(sections) == 5
    assert len(subsections) == 0
    assert sections[0].heading == "Short title."
    assert sections[1].heading == "Commencement."


def test_albury_wodonga_1982_shape2_fixture():
    # Shape 2: fused section+subsection, "1. (1) ...".
    # Expect section 1 containing subsection (1) and subsection (2).
    doc = Document(str(
        CORPUS_DOCX / "albury-wodonga-development-(financial-assistance)-amendment-act-1982-vol0.docx"
    ))
    paras = list(iter_paragraphs(doc))
    sections = [p for p in paras if p.element_type == ElementType.SECTION]
    subsections = [p for p in paras if p.element_type == ElementType.SUBSECTION]
    assert any(s.number == "1" for s in sections)
    subsec_numbers = {s.number for s in subsections}
    assert "1" in subsec_numbers
    assert "2" in subsec_numbers


def test_act_supreme_court_1992_heading_without_style():
    # PART 1—PRELIMINARY styled as plain bold Normal text, no ActHead
    # paragraphs anywhere; tests _HEADING_RE running without the style gate.
    doc = Document(str(
        CORPUS_DOCX / "a.c.t.-supreme-court-(transfer)-act-1992-vol0.docx"
    ))
    paras = list(iter_paragraphs(doc))
    parts = [p for p in paras if p.element_type == ElementType.PART]
    assert len(parts) >= 1
    assert any("PRELIMINARY" in p.heading.upper() for p in parts)


def test_agricultural_chemical_levy_1994_shape3_fixture():
    # Shape 3: style-driven section heading ("Heading 5" style, "N  Heading"
    # text) + "subsection"-styled body/numbered-subsection paragraphs.
    # Expect 4 sections (Short title, Commencement, Imposition, Act does
    # not impose levy on property of a State); sections 3 and 4 each have
    # 2 numbered subsections, sections 1 and 2 have none.
    doc = Document(str(
        CORPUS_DOCX / "agricultural-and-veterinary-chemical-products-levy-imposition-(customs)-act-1994-vol0.docx"
    ))
    paras = list(iter_paragraphs(doc))
    sections = [p for p in paras if p.element_type == ElementType.SECTION]
    subsections = [p for p in paras if p.element_type == ElementType.SUBSECTION]
    assert len(sections) == 4
    assert sections[0].heading == "Short title"
    assert sections[2].heading == "Imposition"
    assert {s.number for s in subsections} >= {"1", "2"}


def test_northern_territory_commonwealth_lands_1980_shape3_fixture():
    # Expect 3 sections (Short title, Commencement, Notification of
    # acquisition of certain interests in land); section 3 has 3 numbered
    # subsections.
    doc = Document(str(
        CORPUS_DOCX / "northern-territory-(commonwealth-lands)-act-1980-vol0.docx"
    ))
    paras = list(iter_paragraphs(doc))
    sections = [p for p in paras if p.element_type == ElementType.SECTION]
    subsections = [p for p in paras if p.element_type == ElementType.SUBSECTION]
    assert len(sections) == 3
    assert sections[0].heading == "Short title"
    assert sections[2].heading == "Notification of acquisition of certain interests in land"
    assert {s.number for s in subsections} >= {"1", "2", "3"}


def test_loan_war_service_land_settlement_1970_shape3_fixture():
    # Expect 4 sections (Short title, Commencement, Authority to borrow
    # $4,500,000, Application of moneys), 0 subsections (all bodies are
    # plain "subsection"-styled prose with no "(N)" numbering).
    doc = Document(str(
        CORPUS_DOCX / "loan-(war-service-land-settlement)-act-1970-vol0.docx"
    ))
    paras = list(iter_paragraphs(doc))
    sections = [p for p in paras if p.element_type == ElementType.SECTION]
    assert len(sections) == 4
    assert sections[0].heading == "Short title"
    assert sections[3].heading == "Application of moneys"
