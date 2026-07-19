#!/usr/bin/env python3
"""Task 1: cluster all Acts with an empty <body> by DOCX style-signature,
then dump a stratified sample of each cluster's raw paragraph
(style, text, all_bold) sequence for manual shape confirmation.

Usage:
    python scripts/discover_legacy_shapes.py [--corpus-dir corpus] [--sample-per-cluster 5]

Output: <corpus-dir>/reports/legacy-shape-discovery.json
"""
from __future__ import annotations

import argparse
import json
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from lxml import etree

AKN_NS = "http://docs.oasis-open.org/legaldocml/ns/akn/3.0"
NS = {"akn": AKN_NS}

_IGNORED_STYLES = {"Normal", "Body Text", "Default"}


def _act_has_empty_body(xml_path: Path) -> bool:
    try:
        root = etree.parse(str(xml_path)).getroot()
    except Exception:
        return False
    return len(root.findall(".//akn:section", NS)) == 0


def _style_signature(doc: Document) -> frozenset[str]:
    sig = set()
    for para in doc.paragraphs:
        name = para.style.name if para.style else "Default"
        if name not in _IGNORED_STYLES:
            sig.add(name)
    return frozenset(sig)


def _all_bold(para) -> bool:
    non_ws = [r for r in para.runs if r.text.strip()]
    return bool(non_ws) and all(bool(r.bold) for r in non_ws)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--corpus-dir", default="corpus", type=Path)
    ap.add_argument("--sample-per-cluster", default=5, type=int)
    args = ap.parse_args()

    xml_dir = args.corpus_dir / "xml"
    docx_dir = args.corpus_dir / "docx"

    legacy_acts: list[tuple[str, Path]] = []
    for xml_path in sorted(xml_dir.glob("*.xml")):
        if not _act_has_empty_body(xml_path):
            continue
        docx_path = docx_dir / f"{xml_path.stem}-vol0.docx"
        if docx_path.exists():
            legacy_acts.append((xml_path.stem, docx_path))

    print(f"Found {len(legacy_acts)} Acts with empty <body> (expect ~550)")

    clusters: dict[frozenset[str], list[str]] = defaultdict(list)
    for name, docx_path in legacy_acts:
        doc = Document(str(docx_path))
        clusters[_style_signature(doc)].append(name)

    by_size = sorted(clusters.items(), key=lambda kv: -len(kv[1]))

    report = {"total_legacy_acts": len(legacy_acts), "clusters": []}
    for sig, names in by_size:
        sample_names = names[: args.sample_per_cluster]
        sample_paras = []
        for name in sample_names:
            doc = Document(str(docx_dir / f"{name}-vol0.docx"))
            paras = [
                {
                    "style": p.style.name if p.style else "Default",
                    "text": p.text.strip(),
                    "all_bold": _all_bold(p),
                }
                for p in doc.paragraphs
                if p.text.strip()
            ][:30]
            sample_paras.append({"act": name, "paragraphs": paras})
        report["clusters"].append({
            "style_signature": sorted(sig) or ["(no distinguishing style)"],
            "act_count": len(names),
            "sample_acts": sample_names,
            "sample_paragraphs": sample_paras,
        })

    out_path = args.corpus_dir / "reports" / "legacy-shape-discovery.json"
    out_path.write_text(json.dumps(report, ensure_ascii=False, indent=2))

    print(f"\n{len(clusters)} distinct style-signature clusters:")
    for sig, names in by_size:
        label = ", ".join(sorted(sig)[:4]) or "(no distinguishing style)"
        print(f"  {len(names):>4}  {label}")
    print(f"\nFull sample dump written to {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
