from __future__ import annotations

import json
import sys
from dataclasses import asdict
from pathlib import Path

import click
from docx import Document
from huggingface_hub import HfApi

from lexau.corpus import Corpus
from lexau.crawler import Crawler
from lexau.builder import AknBuilder
from lexau.docx_reader import iter_paragraphs


def _find_endnote_volume(docx_paths: list[Path]) -> Path | None:
    """Return the volume containing 'Endnotes' (ENotesHeading 1), scanning from last."""
    for path in reversed(docx_paths):
        doc = Document(str(path))
        for para in doc.paragraphs:
            if para.style and para.style.name == "ENotesHeading 1" and para.text.strip() == "Endnotes":
                return path
    return None


@click.group()
def cli() -> None:
    """lex-au: Commonwealth Acts as AKN 3.0 XML."""


def _build_acts(act_names: list[str], corpus_dir: Path, force: bool, doc_type: str = "act") -> None:
    """Build Acts: download DOCX and convert to AKN XML."""
    corpus = Corpus(corpus_dir)
    crawler = Crawler()
    docx_dir = corpus_dir / "docx"
    reports_dir = corpus_dir / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)

    index_path = corpus_dir / "index.json"
    corpus_index: dict = {}
    if index_path.exists():
        raw = json.loads(index_path.read_text())
        corpus_index = {
            entry["name"]: {"frbr_uri": entry.get("frbr_uri", "")}
            for entry in raw.get("acts", {}).values()
        }

    report_rows: list = []

    for act_name in act_names:
        try:
            click.echo(f"[fetch ] {act_name}")
            meta = crawler.fetch_metadata(act_name, doc_type=doc_type)
            if meta is None:
                click.echo(f"  SKIP -- not found in legislation.gov.au", err=True)
                continue

            if not force and corpus.is_current(meta):
                click.echo(f"  SKIP -- compilation #{meta.comp_num} already current")
                continue

            docx_paths = crawler.fetch_docx_volumes(meta, docx_dir)
            if not docx_paths:
                click.echo(f"  SKIP -- DOCX download failed", err=True)
                continue

            click.echo(f"[convert] {act_name} ({len(docx_paths)} volume(s))")
            builder = AknBuilder(meta)
            for docx_path in docx_paths:
                doc = Document(docx_path)
                for p in iter_paragraphs(doc):
                    builder.add(p)

            endnote_vol = _find_endnote_volume(docx_paths)
            xml, report = builder.build_with_report(corpus_index, last_volume_path=endnote_vol)
            report.volumes_fetched = len(docx_paths)

            saved = corpus.save(meta, xml)
            click.echo(f"  saved -> {saved.relative_to(corpus_dir)}")

            report_path = reports_dir / f"{meta.safe_name}-v0.5.0.json"
            report_path.write_text(json.dumps(asdict(report), ensure_ascii=False, indent=2))
            report_rows.append(report)

        except Exception as exc:  # noqa: BLE001
            click.echo(f"  ERROR -- {act_name}: {exc}", err=True)
            continue

    if report_rows:
        click.echo("\n--- Parse Report Summary ---")
        click.echo(
            f"{'Act':<40} {'Vols':>4} {'Subsecs':>7} {'Paras':>6} "
            f"{'Sched':>5} {'Claus':>5} {'Notes':>5} {'Pens':>4} {'Tabs':>4} "
            f"{'Refs':>5} {'Unres':>5} {'Fallbk':>6} "
            f"{'Terms':>5} {'DupT':>4} {'Qtys':>4} {'Roles':>5} {'NtRef':>5}"
        )
        click.echo("-" * 130)
        for r in report_rows:
            click.echo(
                f"{r.act_name:<40} {r.volumes_fetched:>4} {r.subsections_parsed:>7} "
                f"{r.paragraphs_parsed:>6} {r.schedules_found:>5} {r.schedule_clauses_found:>5} "
                f"{r.notes_found:>5} {r.penalties_found:>4} {r.tables_found:>4} "
                f"{r.refs_resolved:>5} {r.refs_unresolved:>5} {r.style_fallbacks:>6} "
                f"{r.terms_found:>5} {r.duplicate_terms:>4} {r.quantities_found:>4} "
                f"{r.roles_found:>5} {r.note_refs_injected:>5}"
            )

    click.echo("\nDone.")


@cli.command()
@click.option("--acts", multiple=True, help="Act name(s) to build (repeatable)")
@click.option(
    "--list-file",
    type=click.Path(exists=True, path_type=Path),
    help="File with one Act name per line",
)
@click.option("--all", "build_all", is_flag=True, help="Build all Acts in acts.txt")
@click.option(
    "--corpus-dir",
    type=click.Path(path_type=Path),
    default=Path("corpus"),
    show_default=True,
)
@click.option("--force", is_flag=True, help="Re-convert even if compilation is current")
@click.option(
    "--type", "doc_type",
    default="act",
    type=click.Choice(["act", "regulation", "instrument"]),
    show_default=True,
    help="Document type (act, regulation, or instrument)",
)
def build(
    acts: tuple[str, ...],
    list_file: Path | None,
    build_all: bool,
    corpus_dir: Path,
    force: bool,
    doc_type: str,
) -> None:
    """Download DOCX and convert to AKN XML."""
    act_names = list(acts)

    if list_file:
        act_names += [ln.strip() for ln in list_file.read_text().splitlines() if ln.strip()]

    if build_all or (not act_names and not list_file):
        default_list = Path("acts.txt")
        if not default_list.exists():
            click.echo("No acts specified and acts.txt not found.", err=True)
            sys.exit(1)
        act_names += [ln.strip() for ln in default_list.read_text().splitlines() if ln.strip()]

    _build_acts(act_names, corpus_dir, force, doc_type=doc_type)


@cli.command()
@click.option("--corpus-dir", type=click.Path(path_type=Path), default=Path("corpus"), show_default=True)
@click.option("--site-dir", type=click.Path(path_type=Path), default=Path("site"), show_default=True)
@click.option("--templates-dir", type=click.Path(path_type=Path), default=Path("templates"), show_default=True)
def site(corpus_dir: Path, site_dir: Path, templates_dir: Path) -> None:
    """Generate static HTML site from the corpus."""
    from lexau.site import SiteGenerator
    corpus = Corpus(corpus_dir)
    gen = SiteGenerator(corpus, site_dir, templates_dir)
    gen.generate()
    click.echo(f"Site generated -> {site_dir}/")


@cli.command("export-hf")
@click.option("--repo", required=True, help="HF dataset repo, e.g. cchew/lex-au")
@click.option("--corpus-dir", type=click.Path(path_type=Path), default=Path("corpus"), show_default=True)
@click.option(
    "--readme",
    type=click.Path(exists=True, path_type=Path),
    default=Path("hf-readme.md"),
    show_default=True,
    help="Dataset card to upload as README.md",
)
def export_hf(repo: str, corpus_dir: Path, readme: Path) -> None:
    """Push corpus XML + index + dataset card to a Hugging Face dataset."""
    api = HfApi()
    click.echo(f"Uploading corpus to {repo}…")
    api.upload_folder(
        folder_path=str(corpus_dir),
        repo_id=repo,
        repo_type="dataset",
        commit_message="lex-au corpus update",
        ignore_patterns=["docx/**"],
    )
    click.echo(f"Uploading dataset card from {readme}…")
    api.upload_file(
        path_or_fileobj=str(readme),
        path_in_repo="README.md",
        repo_id=repo,
        repo_type="dataset",
        commit_message="lex-au dataset card update",
    )
    click.echo("Upload complete.")


@cli.command("export-jsonl")
@click.option("--corpus-dir", type=click.Path(path_type=Path), default=Path("corpus"), show_default=True)
def export_jsonl(corpus_dir: Path) -> None:
    """Write corpus/data/train.jsonl — one row per Act, from index.json.

    HF's auto-parquet-conversion / Data Viewer can't parse .xml (skipped
    entirely) or a single-dict index.json (not row-shaped), and falls back to
    inferring a schema from reports/*.json, which drifts across pipeline
    versions. This gives it a proper tabular file to key off instead —
    referenced explicitly via hf-readme.md's `configs:` block.
    """
    index_path = corpus_dir / "index.json"
    index = json.loads(index_path.read_text())
    data_dir = corpus_dir / "data"
    data_dir.mkdir(parents=True, exist_ok=True)
    out_path = data_dir / "train.jsonl"

    rows = 0
    with out_path.open("w") as f:
        for slug, entry in sorted(index["acts"].items()):
            row = {
                "slug": slug,
                "name": entry["name"],
                "title_id": entry["title_id"],
                "comp_id": entry["comp_id"],
                "comp_num": entry["comp_num"],
                "year": entry["year"],
                "number": entry["number"],
                "effective_date": entry["effective_date"],
                "xml_path": entry["xml_path"],
            }
            f.write(json.dumps(row) + "\n")
            rows += 1

    click.echo(f"Wrote {rows} rows -> {out_path}")


@cli.command("list-acts")
@click.option(
    "--output",
    "-o",
    type=click.Path(path_type=Path),
    default=None,
    help="Write to file instead of stdout (e.g. acts.txt)",
)
@click.option("--page-size", default=200, show_default=True, help="OData page size per request")
def list_acts(output: Path | None, page_size: int) -> None:
    """List all in-force Commonwealth Acts from legislation.gov.au."""
    crawler = Crawler()
    click.echo("Fetching Act list from legislation.gov.au…", err=True)
    names = crawler.list_acts(page_size=page_size)
    click.echo(f"Found {len(names)} Acts.", err=True)
    text = "\n".join(names) + "\n"
    if output:
        output.write_text(text)
        click.echo(f"Written to {output}", err=True)
    else:
        click.echo(text, nl=False)


@cli.command("list-instruments")
@click.option(
    "--output",
    "-o",
    type=click.Path(path_type=Path),
    default=None,
    help="Write to file instead of stdout (e.g. instruments.txt)",
)
@click.option("--page-size", default=200, show_default=True, help="OData page size per request")
def list_instruments(output: Path | None, page_size: int) -> None:
    """List all in-force Commonwealth legislative instruments from legislation.gov.au."""
    crawler = Crawler()
    click.echo("Fetching instrument list from legislation.gov.au…", err=True)
    names = crawler.list_instruments(page_size=page_size)
    click.echo(f"Found {len(names)} instruments.", err=True)
    text = "\n".join(names) + "\n"
    if output:
        output.write_text(text)
        click.echo(f"Written to {output}", err=True)
    else:
        click.echo(text, nl=False)


@cli.command()
@click.option("--since", required=True, help="ISO date, e.g. 2026-01-01")
@click.option(
    "--corpus-dir",
    type=click.Path(path_type=Path),
    default=Path("corpus"),
    show_default=True,
)
def update(since: str, corpus_dir: Path) -> None:
    """Fetch and re-convert Acts modified since a given date."""
    from datetime import date as _date

    since_date = _date.fromisoformat(since)

    crawler = Crawler()
    click.echo(f"Checking for Acts modified since {since}...")
    modified = crawler.list_modified_since(since_date)

    if not modified:
        click.echo("No modified Acts found.")
        return

    # Only re-build Acts already in the corpus; legislation.gov.au returns
    # modifications across all Acts, not just our corpus members.
    corpus = Corpus(corpus_dir)
    corpus_names = {meta.name for meta in corpus.all_metadata()}
    in_corpus = [name for name in modified if name in corpus_names]

    if not in_corpus:
        click.echo(
            f"Found {len(modified)} modified Act(s), none in the corpus."
        )
        return

    click.echo(f"Found {len(in_corpus)} modified corpus Act(s): {', '.join(in_corpus)}")
    _build_acts(in_corpus, corpus_dir, force=True)
